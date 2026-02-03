import os
import logging
from typing import Dict, Any, Optional, List, Tuple
from io import BytesIO
from PIL import Image
import pytesseract
import PyPDF2
from docx import Document
import ollama

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import configuration constants
from app.constants import (
    VISION_MODEL_NAME,
    VISION_TEMPERATURE_CLASSIFY,
    VISION_TEMPERATURE_EXTRACT,
    VISION_PREVIEW_LENGTH,
    IMAGE_TYPE_TEXT_SCREENSHOT,
    IMAGE_TYPE_COMPLEX,
    IMAGE_TYPE_UNKNOWN,
    EXTRACTION_METHOD_OCR_FAST,
    EXTRACTION_METHOD_VISION_AI,
    EXTRACTION_METHOD_OCR_FALLBACK,
    FILE_TYPE_IMAGE,
    FILE_TYPE_PDF,
    FILE_TYPE_DOCX,
    FILE_TYPE_TEXT,
    SUPPORTED_IMAGE_EXTENSIONS,
    SUPPORTED_PDF_EXTENSIONS,
    SUPPORTED_DOCX_EXTENSIONS,
    SUPPORTED_TEXT_EXTENSIONS,
    TEXT_ENCODING_UTF8,
    TEXT_ENCODING_LATIN1,
    PROMPT_CLASSIFY_IMAGE,
    PROMPT_EXTRACT_TEXT,
    PDF_PAGE_SEPARATOR,
    DOCX_LINE_SEPARATOR,
    DOCX_TABLE_CELL_SEPARATOR,
    ERROR_UNSUPPORTED_FILE_TYPE,
    ERROR_HANDLER_NOT_IMPLEMENTED,
    ERROR_FILE_PROCESSING_FAILED,
    ERROR_IMAGE_PROCESSING_FAILED,
    ERROR_PDF_EXTRACTION_FAILED,
    ERROR_DOCX_EXTRACTION_FAILED,
    ERROR_TEXT_EXTRACTION_FAILED,
    ERROR_BOTH_METHODS_FAILED,
    LOG_VISION_MODEL_NOT_FOUND,
    LOG_VISION_CONNECTION_ERROR,
    LOG_CLASSIFYING_IMAGE,
    LOG_IMAGE_TYPE_TEXT,
    LOG_IMAGE_TYPE_COMPLEX,
    LOG_CLASSIFICATION_ERROR,
    LOG_EXTRACTING_TEXT,
    LOG_EXTRACTED_TEXT_PREFIX,
    LOG_EMPTY_RESPONSE,
    LOG_VISION_ERROR,
    LOG_USING_OCR_FAST,
    LOG_USING_VISION,
    LOG_OCR_FAILED_FALLBACK,
    LOG_VISION_UNAVAILABLE,
    WARNING_VISION_MODEL_UNAVAILABLE
)

# Import profiler utilities
from app.utils.profiler import timed_operation, track_performance

# Import helper functions
from app.utils.document_helpers import (
    create_error_response,
    create_success_response,
    ensure_rgb_image,
    save_temp_image,
    cleanup_temp_file,
    extract_model_name
)

class DocumentProcessor:
    """
    Process various document formats and extract text
    
    Supports: images (PNG, JPG, etc.), PDF, DOCX, and plain text files
    Uses vision AI (MiniCPM-V) for intelligent image text extraction
    """
    
    def __init__(self):
        self.supported_formats: Dict[str, List[str]] = {
            FILE_TYPE_IMAGE: SUPPORTED_IMAGE_EXTENSIONS,
            FILE_TYPE_PDF: SUPPORTED_PDF_EXTENSIONS,
            FILE_TYPE_DOCX: SUPPORTED_DOCX_EXTENSIONS,
            FILE_TYPE_TEXT: SUPPORTED_TEXT_EXTENSIONS
        }
        self.vision_model: str = VISION_MODEL_NAME
        # Cache vision model availability to avoid repeated checks
        self._vision_model_available: Optional[bool] = None
    
    def get_file_type(self, filename: str) -> Optional[str]:
        """
        Determine file type from extension
        
        Args:
            filename: Name of file to check
            
        Returns:
            File type string or None if unsupported
        """
        ext = os.path.splitext(filename)[1].lower()
        for file_type, extensions in self.supported_formats.items():
            if ext in extensions:
                return file_type
        return None
    
    def process_file(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """
        Process uploaded file and extract text content
        
        Routes to appropriate handler based on file type
        
        Args:
            file_data: Raw file bytes
            filename: Original filename
            
        Returns:
            Dict with extracted text and metadata, or error information
        """
        file_type = self.get_file_type(filename)
        
        if not file_type:
            error_msg = f'{ERROR_UNSUPPORTED_FILE_TYPE}: {os.path.splitext(filename)[1]}'
            return create_error_response(filename, error_msg)
        
        try:
            # Route to appropriate handler based on file type
            handlers = {
                FILE_TYPE_IMAGE: self._process_image,
                FILE_TYPE_PDF: self._process_pdf,
                FILE_TYPE_DOCX: self._process_docx,
                FILE_TYPE_TEXT: self._process_text
            }
            
            handler = handlers.get(file_type)
            if handler:
                return handler(file_data, filename)
            else:
                error_msg = f'{ERROR_HANDLER_NOT_IMPLEMENTED} {file_type}'
                return create_error_response(filename, error_msg)
                
        except Exception as e:
            logger.exception(f"Error processing {filename}")
            error_msg = f'{ERROR_FILE_PROCESSING_FAILED}: {str(e)}'
            return create_error_response(filename, error_msg)
    
    def _check_vision_model(self) -> bool:
        """
        Check if Ollama vision model is available
        
        Uses caching to avoid repeated API calls
        
        Returns:
            True if vision model is available, False otherwise
        """
        # Return cached result if available
        if self._vision_model_available is not None:
            return self._vision_model_available
        
        try:
            models_response = ollama.list()
            
            # Handle both dict and list response formats
            if isinstance(models_response, dict):
                models_list = models_response.get('models', [])
            else:
                models_list = models_response
            
            # Check if our model is in the list
            for model in models_list:
                model_name = extract_model_name(model)
                if self.vision_model in model_name:
                    self._vision_model_available = True
                    return True
            
            logger.warning(
                f"[Vision Model] {self.vision_model} {LOG_VISION_MODEL_NOT_FOUND} {self.vision_model}"
            )
            self._vision_model_available = False
            return False
            
        except Exception as e:
            logger.error(f"[Vision Model] {LOG_VISION_CONNECTION_ERROR}: {e}")
            self._vision_model_available = False
            return False
    
    def _call_vision_model(
        self,
        image_path: str,
        prompt: str,
        temperature: float
    ) -> Optional[str]:
        """
        Call Ollama vision model with image and prompt
        
        Centralized vision model API call to eliminate duplication
        
        Args:
            image_path: Path to temporary image file
            prompt: Prompt for the vision model
            temperature: Temperature setting for model
            
        Returns:
            Model response text or None on error
        """
        try:
            response = ollama.chat(
                model=self.vision_model,
                messages=[
                    {
                        'role': 'user',
                        'content': prompt,
                        'images': [image_path]
                    }
                ],
                options={
                    'temperature': temperature
                }
            )
            return response['message']['content']
        except Exception as e:
            logger.error(f"[{LOG_VISION_ERROR}]: {e}")
            return None

    def _classify_image_type(self, image: Image.Image) -> str:
        """
        Use vision model to determine if image is text screenshot or complex image
        
        Args:
            image: PIL Image to classify
            
        Returns:
            IMAGE_TYPE_TEXT_SCREENSHOT, IMAGE_TYPE_COMPLEX, or IMAGE_TYPE_UNKNOWN
        """
        if not self._check_vision_model():
            return IMAGE_TYPE_UNKNOWN
        
        tmp_path = None
        try:
            # Ensure RGB format
            image = ensure_rgb_image(image)
            
            # Save to temporary file for Ollama
            tmp_path = save_temp_image(image)
            
            logger.info(LOG_CLASSIFYING_IMAGE)
            
            # Call vision model for classification
            classification = self._call_vision_model(
                tmp_path,
                PROMPT_CLASSIFY_IMAGE,
                VISION_TEMPERATURE_CLASSIFY
            )
            
            if not classification:
                return IMAGE_TYPE_UNKNOWN
            
            classification = classification.strip().lower()
            
            if 'yes' in classification:
                logger.info(LOG_IMAGE_TYPE_TEXT)
                return IMAGE_TYPE_TEXT_SCREENSHOT
            else:
                logger.info(LOG_IMAGE_TYPE_COMPLEX)
                return IMAGE_TYPE_COMPLEX
                
        except Exception as e:
            logger.error(f"{LOG_CLASSIFICATION_ERROR}: {e}")
            return IMAGE_TYPE_UNKNOWN
        finally:
            # Clean up temp file
            if tmp_path:
                cleanup_temp_file(tmp_path)
    
    def _extract_text_with_vision(self, image: Image.Image) -> Optional[str]:
        """
        Use Ollama vision model to extract text and description from image
        
        Args:
            image: PIL Image to extract text from
            
        Returns:
            Extracted text and description, or None on error
        """
        if not self._check_vision_model():
            return None
        
        tmp_path = None
        try:
            # Ensure RGB format
            image = ensure_rgb_image(image)
            
            # Save to temporary file for Ollama
            tmp_path = save_temp_image(image)
            
            logger.info(f"{LOG_EXTRACTING_TEXT} {self.vision_model}...")
            
            # Call vision model for text extraction
            extracted_text = self._call_vision_model(
                tmp_path,
                PROMPT_EXTRACT_TEXT,
                VISION_TEMPERATURE_EXTRACT
            )
            
            if extracted_text:
                logger.info(f"{LOG_EXTRACTED_TEXT_PREFIX} {extracted_text[:100]}...")
                return extracted_text
            else:
                logger.warning(LOG_EMPTY_RESPONSE)
                return None
                
        except Exception as e:
            logger.error(f"[{LOG_VISION_ERROR}]: {e}")
            return None
        finally:
            # Clean up temp file
            if tmp_path:
                cleanup_temp_file(tmp_path)
    
    @timed_operation("process_image")
    def _process_image(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """
        Smart image text extraction with vision model classification
        
        Strategy:
        1. Classify image type (text screenshot vs complex image)
        2. Text screenshots: Fast pytesseract OCR
        3. Complex images: Accurate vision model extraction
        4. Fallback: pytesseract OCR if vision unavailable
        
        Args:
            file_data: Raw image bytes
            filename: Original filename
            
        Returns:
            Processing result with text and metadata
        """
        try:
            image = Image.open(BytesIO(file_data))
            
            # Step 1: Classify image type with vision model
            image_type = self._classify_image_type(image)
            
            # Step 2: Route to appropriate extraction method
            if image_type == IMAGE_TYPE_TEXT_SCREENSHOT:
                # Fast OCR for simple text screenshots
                logger.info(LOG_USING_OCR_FAST)
                try:
                    ocr_text = pytesseract.image_to_string(image).strip()
                    word_count = len(ocr_text.split())
                    
                    logger.debug(ocr_text)
                    return create_success_response(
                        text=ocr_text,
                        filename=filename,
                        file_type=FILE_TYPE_IMAGE,
                        extraction_method=EXTRACTION_METHOD_OCR_FAST,
                        metadata={
                            'size': image.size,
                            'mode': image.mode,
                            'format': image.format,
                            'word_count': word_count,
                            'image_type': IMAGE_TYPE_TEXT_SCREENSHOT
                        }
                    )
                except Exception as ocr_error:
                    logger.warning(f"{LOG_OCR_FAILED_FALLBACK}: {ocr_error}")
                    # Fall through to vision model
            
            # Use vision model for complex images or if OCR failed
            logger.info(LOG_USING_VISION)
            vision_text = self._extract_text_with_vision(image)
            
            logger.debug(vision_text)
            if vision_text:
                return create_success_response(
                    text=vision_text,
                    filename=filename,
                    file_type=FILE_TYPE_IMAGE,
                    extraction_method=EXTRACTION_METHOD_VISION_AI,
                    metadata={
                        'size': image.size,
                        'mode': image.mode,
                        'format': image.format,
                        'model': self.vision_model,
                        'image_type': image_type
                    }
                )
            
            # Fallback to pytesseract OCR if vision model unavailable
            logger.warning(LOG_VISION_UNAVAILABLE)
            try:
                ocr_text = pytesseract.image_to_string(image).strip()
                word_count = len(ocr_text.split())
                
                return create_success_response(
                    text=ocr_text,
                    filename=filename,
                    file_type=FILE_TYPE_IMAGE,
                    extraction_method=EXTRACTION_METHOD_OCR_FALLBACK,
                    metadata={
                        'size': image.size,
                        'mode': image.mode,
                        'format': image.format,
                        'word_count': word_count,
                        'warning': WARNING_VISION_MODEL_UNAVAILABLE
                    }
                )
            except Exception as ocr_error:
                return create_error_response(
                    f"{ERROR_BOTH_METHODS_FAILED}: {str(ocr_error)}",
                    filename
                )
            
        except Exception as e:
            return create_error_response(
                f"{ERROR_IMAGE_PROCESSING_FAILED}: {str(e)}",
                filename
            )
    
    @timed_operation("process_pdf")
    def _process_pdf(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from PDF document
        
        Args:
            file_data: Raw PDF bytes
            filename: Original filename
            
        Returns:
            Processing result with text and metadata
        """
        try:
            pdf_file = BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            full_text = PDF_PAGE_SEPARATOR.join(text_parts)
            
            return create_success_response(
                text=full_text.strip(),
                filename=filename,
                file_type=FILE_TYPE_PDF,
                metadata={
                    'num_pages': len(pdf_reader.pages),
                    'metadata': pdf_reader.metadata
                }
            )
        except Exception as e:
            return create_error_response(
                f"{ERROR_PDF_EXTRACTION_FAILED}: {str(e)}",
                filename
            )
    
    @timed_operation("process_docx")
    def _process_docx(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from DOCX document including tables
        
        Args:
            file_data: Raw DOCX bytes
            filename: Original filename
            
        Returns:
            Processing result with text and metadata
        """
        try:
            docx_file = BytesIO(file_data)
            doc = Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = DOCX_TABLE_CELL_SEPARATOR.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = DOCX_LINE_SEPARATOR.join(text_parts)
            
            return create_success_response(
                text=full_text.strip(),
                filename=filename,
                file_type=FILE_TYPE_DOCX,
                metadata={
                    'num_paragraphs': len(doc.paragraphs),
                    'num_tables': len(doc.tables)
                }
            )
        except Exception as e:
            return create_error_response(
                f"{ERROR_DOCX_EXTRACTION_FAILED}: {str(e)}",
                filename
            )
    
    @timed_operation("process_text")
    def _process_text(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from plain text file with encoding fallback
        
        Args:
            file_data: Raw text file bytes
            filename: Original filename
            
        Returns:
            Processing result with text and metadata
        """
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                text = file_data.decode(TEXT_ENCODING_UTF8)
            except UnicodeDecodeError:
                text = file_data.decode(TEXT_ENCODING_LATIN1)
            
            return create_success_response(
                text=text.strip(),
                filename=filename,
                file_type=FILE_TYPE_TEXT,
                metadata={
                    'size_bytes': len(file_data)
                }
            )
        except Exception as e:
            return create_error_response(
                f"{ERROR_TEXT_EXTRACTION_FAILED}: {str(e)}",
                filename
            )
    
    def is_supported(self, filename: str) -> bool:
        """Check if file type is supported"""
        return self.get_file_type(filename) is not None
    
    def get_supported_extensions(self) -> list:
        """Get list of all supported extensions"""
        extensions = []
        for exts in self.supported_formats.values():
            extensions.extend(exts)
        return extensions


# Singleton instance
_processor_instance = None

def get_document_processor() -> DocumentProcessor:
    """Get or create document processor instance"""
    global _processor_instance
    if _processor_instance is None:
        _processor_instance = DocumentProcessor()
    return _processor_instance
